import json
import os

from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel
from pypdf import PdfReader
from docx import Document

load_dotenv()
my_api_key = os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("API KEY NOT PRESENT")

client = Groq(api_key=my_api_key)

model = "openai/gpt-oss-120b"

class Job_D(BaseModel):
    role:str
    required_skills:list[str]
    preferred_skills: list[str]
    minimum_experience :float | None
    education_requirements:list[str]

def parse_job_description(job_description):

    jobd_schema = Job_D.model_json_schema()

    system_prompt = f"""
    You are an expert HR assistant.
    Your task is to analyze job descriptions and extract
    structured information from them.

    Return only valid JSON matching this schema:
    {jobd_schema}

    Important:
    Do not return schema itself.
    Do not return fields like "properties", "title" or "type".
    Fill the schema with actual information extracted from job description.

    If minimum experience is not mentioned return null.
    If information for a list is missing return an empty list.
    Do not invent information.
    """

    user_prompt = f"""
    Analyze the following job description:

    {job_description}
    """

    message_system = {
        "role": "system",
        "content": system_prompt
    }

    message_user = {
        "role": "user",
        "content": user_prompt
    }

    messages = [message_system, message_user]

    response = client.chat.completions.create(
        model=model,
        messages=messages,
        response_format={"type": "json_object"}
    )

    raw_json = response.choices[0].message.content

    job_data = json.loads(raw_json)

    job = Job_D(**job_data)

    return job

class MatchResult(BaseModel):
    score:float
    details: dict
class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []

class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None

    total_experience_years: float | None = None

    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []

resume_schema = Resume.model_json_schema()
def final_score(job,resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {match_schema}

    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """
    message={
        "role": "user",
        "content" : prompt
    }
    messages=[message]
    response_format={
        "type": "json_object"
    }
    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)

def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }
    response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume

from pypdf import PdfReader
from docx import Document


def read_pdf(file_path):
    try:
        reader = PdfReader(file_path)
        text = ""

        for page in reader.pages:
            page_text = page.extract_text()

            if page_text:
                text += page_text + "\n"

        return text
    except Exception as error:
        raise ValueError("The uploaded PDF could not be read.") from error


def read_docx(file_path):
    try:
        document = Document(file_path)
        text = ""

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"

        return text
    except Exception as error:
        raise ValueError("The uploaded DOCX file could not be read.") from error


def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)

    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)

    else:
        return None
